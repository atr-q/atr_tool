import duckdb, pandas
import math
import os, sys, errno
import subprocess as sp
import docx
from docx.shared import Pt
from docx.shared import Cm
import customtkinter
from PIL import Image
from tkinter import filedialog

customtkinter.set_appearance_mode("dark")
root = customtkinter.CTk()
root.minsize(width=1280, height=720)
root.after(0, lambda:root.state('zoomed'))
root.title('Whatnot Seller Tool - By Abandoned Treasures Reclaimed')
pf = os.getenv("ProgramFiles")
microEditor = pf + "\\Microsoft Office\\root\\Office16\\WINWORD.EXE"
wordEditor = pf + "\\Windows NT\\Accessories\\wordpad.exe"
defEditor = "notepad.exe"
radio_var = customtkinter.StringVar(value="n")


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

root.iconbitmap(resource_path('16_3x_XIg_icon.ico'))

def silent_remove(filename):
    try:
        os.remove(resource_path(filename))
    except OSError as e: # this would be "except OSError, e:" before Python 2.6
        if e.errno != errno.ENOENT: # errno.ENOENT = no such file or directory
            raise # re-raise exception if a different error occurred

def get_all_names(np_arr):
    unique_names = []
    for i in range(0, len(np_arr)):
        unique_names.append(np_arr[i][2])
    unique_names = list(set(unique_names))
    unique_names.sort()
    return unique_names

def get_all_categories(np_arr):
    unique_categories = []
    for i in range(0, len(np_arr)):
        unique_categories.append(str(np_arr[i][22]))
    unique_categories = list(set(unique_categories))
    unique_categories.sort()
    return unique_categories

def get_all_coupons(np_arr):
    unique_coupons = []
    for i in range(0, len(np_arr)):
        if str(np_arr[i][16]) != 'nan':
            unique_coupons.append(np_arr[i][16])
    unique_coupons = list(set(unique_coupons))
    return unique_coupons

def rename_columns(livestream_report):
    df = pandas.read_csv(livestream_report)
    df = df.rename(columns=({'buyer': 'Buyer'}))
    df = df.rename(columns=({'product name': 'product_name'}))
    df = df.rename(columns=({'cancelled or failed': 'isCanceled'}))
    df = df.rename(columns=({'sold price': 'sold_price'}))
    df = df.rename(columns=({'coupon code': 'coupon_code'}))
    df = df.rename(columns=({'coupon price': 'coupon_price'}))
    df = df.rename(columns=({'gifted to': 'gifted_to'}))
    df = df.rename(columns=({'product quantity': 'product_quantity'}))
    df[['Category', 'Number']] = df.product_name.str.split("#", expand=True)
    df.to_csv('temporary.csv', index=False)
    return df.to_numpy()

def get_coupon_data(code):
    coupon_data = []
    for i in range(0, len(code)):
        try:
            query = ("SELECT coupon_price "
                     + "FROM 'temporary.csv' WHERE isCanceled IS NULL AND coupon_code ='" + str(code[i]) + "'")
            coupon = duckdb.sql(query).df()
            total_discount = 0
            for j in range(0, len(coupon)):
                price = str(coupon.iat[j, 0]).replace('$', '')
                total_discount += float(price)
            coupon_data.append([code[i], total_discount, int(len(coupon))])
        except duckdb.duckdb.IOException:
            print("ERROR: File is open in other program")
        except duckdb.duckdb.ParserException:
            print("Error: Incorrect Parameter specified")
    grand_total = 0
    grand_uses = 0
    for i in range(0, len(coupon_data)):
        grand_total += round(coupon_data[i][1], 2)
        grand_uses += coupon_data[i][2]
    coupon_data.append(['Grand Total', format(grand_total, '.2f'), grand_uses])
    return coupon_data

def get_coupon_str(data):
    coupon_str = ""
    for i in range(0, (len(data)-1)):
        code = data[i][0]
        discount = format(data[i][1], '.2f')
        number = data[i][2]
        if radio_var.get() == 'w' or radio_var.get() == 'm':
            coupon_str += 'Coupon Code: \t\t' + code + '\nDiscount Total: \t\t$' + str(discount) + '\n# of Uses: \t\t' + str(number) + '\n\n'
        elif radio_var.get() == 'n':
            coupon_str += 'Coupon Code: \t\t' + code + '\nDiscount Total: \t$' + str(discount) + '\n# of Uses: \t\t' + str(number) + '\n\n'
    coupon_str += 'Discount Grand Total: \t$' + str(data[len(data)-1][1]) + '\n'
    coupon_str += '# of Coupons Applied: \t' + str(data[len(data)-1][2]) +'\n\n'
    return coupon_str

def get_shipment(name, category):
    cancelled_str = '\nCancelled:\t\t'
    try:                 # Can add Buyer, Category, sold_price, product_name, isCancelled
        query = ("SELECT Number, isCanceled, product_quantity "
                 + "FROM 'temporary.csv' WHERE Category ='" + category + "' AND Buyer='" + name + "' ORDER BY Category")
        shipment = duckdb.sql(query).df()
        if not shipment.empty:
            unsort_shipment = []
            cnt = 1
            for i in range(0, len(shipment)):
                if shipment.iat[i, 1] == None:
                    if math.isnan(shipment.iat[i, 0]):
                        if int(shipment.iat[i, 2]) > 1:
                            unsort_shipment.append(('x' + str(shipment.iat[i, 2])))
                        else:
                            unsort_shipment.append(cnt)
                            cnt += 1

                    else:
                        unsort_shipment.append(int(shipment.iat[i, 0]))
                else:
                    dup_check = 0
                    for j in range(0, len(unsort_shipment)):
                        if str(unsort_shipment[j]) == str(shipment.iat[i, 0]):
                            dup_check = 1
                    if dup_check == 0:
                        cancelled_str += (str(shipment.iat[i, 0]) + "     ")
            try:
                unsort_shipment.sort()
            except TypeError:
                print('a')

            sorted_shipment = unsort_shipment
            str_shipment = ""
            for i in range(0, len(sorted_shipment)):
                if i == 0:
                    str_shipment += str(sorted_shipment[i])
                else:
                    str_shipment += ("     " + str(sorted_shipment[i]))
            if len(cancelled_str) > 15:
                str_shipment += cancelled_str
            return str_shipment
        else:
            return -1
    except duckdb.duckdb.IOException:
        print("ERROR: File is open in other program")
    except duckdb.duckdb.ParserException:
        print("Error: Incorrect Parameter specified")
    return

def get_all_shipments(names, categories):
    all_shipments = []
    for i in range(0, len(names)):
        shipments = []
        for j in range(0, len(categories)):
            str_shipment = get_shipment(names[i], categories[j])
            if str_shipment != -1:
                shipments.append([categories[j], str_shipment])
        all_shipments.append([names[i], shipments])
    return all_shipments

def create_word_doc(info, coupon):

    df = pandas.DataFrame(info)
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    doc.add_heading('Coupon and Rewards Program Discount Info')
    doc.add_paragraph('\n' + coupon)
    for i in range(0, len(df)):
        buyer = df.iat[i,0]
        shipment = df.iat[i, 1]
        if radio_var.get() == 'm':
            doc.add_paragraph('_______________________________________________________________________________', style=None)
        heading = str(i+1) + '.___ ' + buyer
        doc.add_heading(heading, level=1)
        table = doc.add_table(rows=0, cols=2)
        for j in range(0, len(shipment)):
            category = shipment[j][0]
            numbers = shipment[j][1]
            cat_cells = table.add_row().cells
            cat_cells[0].text = 'Category: '
            cat_cells[1].text = str(category)
            num_cells = table.add_row().cells
            num_cells[0].text = 'Item #: '
            num_cells[1].text = str(numbers)
        doc.add_paragraph('\n', style=None)
    filename = 'PRINT ME.docx'
    if radio_var.get() == 'w':
        filename = 'PRINT ME!.docx'

    try:
        doc.save(filename)
        if radio_var.get() == 'm':
            sp.Popen([microEditor, filename])
        elif  radio_var.get() == 'w':
            sp.Popen([wordEditor, filename])
        os.remove('temporary.csv')
        button.configure(text="Upload Livestream Report")
    except PermissionError:
        print('Error: Word Editor is open')
        err_message = ''
        if radio_var.get() == 'w':
            err_message = "Error File Already Open\n\nClose Wordpad and\nClick to Re-Upload Livestream Report"
        elif radio_var.get() == 'm':
            err_message = "Error File Already Open\n\nClose Microsoft Word and\nClick to Re-Upload Livestream Report"
        button.configure(text=err_message)
    return


def create_txt_doc(info, coupon):
    df = pandas.DataFrame(info)
    final_str = '--------------------------------------------------------------------------------\n'
    final_str += 'Coupon and Rewards Program Discount Info\n\n'
    final_str += coupon + '\n\n'
    for i in range(0, len(df)):
        buyer = df.iat[i,0]
        shipment = df.iat[i, 1]
        final_str += '--------------------------------------------------------------------------------\n'
        final_str += str(i+1) + '.___ ' + buyer + '\n'
        for j in range(0, len(shipment)):
            category = shipment[j][0]
            numbers = shipment[j][1]
            final_str += ('Category:\t\t' + category + '\n')
            final_str += ('Item #:   \t\t' + numbers + '\n')
            final_str += '\n\n'
    file_path = 'PRINT ME.txt'
    with open(file_path, 'w') as file:
        file.write(final_str)
    sp.Popen([defEditor, file_path])
    return

def upload_file():
    button.configure(text="Creating Your Document\n\nPlease Wait")
    file_path = filedialog.askopenfilename(filetypes=[(".csv Files", "*.csv")])
    button.configure(text="Upload Livestream Report")
    if file_path:
        file_name = file_path
        file = rename_columns(file_name)
        all_names = get_all_names(file)
        all_categories = get_all_categories(file)
        all_coupons = get_all_coupons(file)
        coupon_arr = get_coupon_data(all_coupons)
        coupon_data = get_coupon_str(coupon_arr)
        shipments = get_all_shipments(all_names, all_categories)
        if radio_var.get() == 'm' or radio_var.get() == 'w':
            create_word_doc(shipments, coupon_data)
        elif radio_var.get() == 'n':
            create_txt_doc(shipments, coupon_data)



contain = customtkinter.CTkScrollableFrame(master=root, fg_color='#0E0C07', scrollbar_button_color='#D6B44A')
contain.pack(fill="both", expand=True)


gui = customtkinter.CTkImage(light_image=Image.open(resource_path('gui.png')),
                                  dark_image=Image.open(resource_path('gui.png')),
                                  size=(1280, 520))

gui_label = customtkinter.CTkLabel(contain, text="", image=gui)
gui_label.pack()

doc_choice_notepad = customtkinter.CTkRadioButton(contain, text="Notepad Document", value='n', fg_color='#D6B44A', hover_color='white', variable=radio_var)
doc_choice_notepad.pack(pady=5)

doc_choice_wordpad = customtkinter.CTkRadioButton(contain, text="Wordpad Document", value='w', fg_color='#D6B44A', hover_color='white', variable=radio_var)
doc_choice_wordpad.pack(pady=5)

doc_choice_docx = customtkinter.CTkRadioButton(contain, text="Microsoft Word Document", value='m', fg_color='#D6B44A', hover_color='white', variable=radio_var)
doc_choice_docx.pack(pady=5)

button = customtkinter.CTkButton(master=contain, text="Upload Livestream Report", fg_color='#D6B44A', hover_color='white', text_color='black', command=upload_file)
button.pack(pady=30, ipady=20, ipadx=10)

root.mainloop()